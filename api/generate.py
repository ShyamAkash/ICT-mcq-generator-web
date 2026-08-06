"""
Vercel Python serverless function.

POST /api/generate
Body (JSON): { "topic": str, "model": str, "qtype": "normal" | "statement" | "code" }

Calls the Gemini API to generate a bilingual (English/Sinhala) MCQ, fills the
matching .docx template, and streams the finished .docx straight back in the
HTTP response. Nothing generated is written to disk or any database -
everything happens in memory for the lifetime of a single request.
"""

import io
import json
import os
import re
from datetime import datetime, timezone
from pathlib import Path

from flask import Flask, request, send_file, jsonify
from docx import Document
from google import genai

from UnicodeToLegacy import ConvertToLegacy

app = Flask(__name__)

# ---------------------------------------------------------------------------
# Paths (resolved from this file's location so they work regardless of the
# process's current working directory).
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates"
PROMPT_PATH = PROJECT_ROOT / "prompt.txt"

TEMPLATE_FILES = {
    "normal": TEMPLATES_DIR / "QuestionNormal.docx",
    "statement": TEMPLATES_DIR / "QuestionStatement.docx",
    "code": TEMPLATES_DIR / "QuestionCode.docx",
}

VALID_TYPES = set(TEMPLATE_FILES.keys())


# ---------------------------------------------------------------------------
# docx helpers (ported as-is from the original desktop app)
# ---------------------------------------------------------------------------
def replace_in_paragraph(paragraph, old, new):
    """Replace text while preserving the formatting of the first run."""
    if old not in paragraph.text:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    full_text = full_text.replace(old, new)
    for i, run in enumerate(paragraph.runs):
        run.text = full_text if i == 0 else ""


def find_and_replace(doc, replacements: dict):
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            replace_in_paragraph(paragraph, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        replace_in_paragraph(paragraph, old, new)

    return doc


def extract_json(raw_text: str) -> dict:
    """Parse the model's response as JSON, defensively stripping any
    accidental markdown code fences the model might add despite the prompt
    asking it not to."""
    text = raw_text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"```\s*$", "", text).strip()
    return json.loads(text)


def build_replacements(resdict: dict, qtype: str) -> dict:
    if qtype == "normal":
        return {
            "Question English": resdict["QEng"],
            "Question Sinhala": ConvertToLegacy(resdict["QSin"]),

            "Answer 1 English": resdict["AnswersEng"][0],
            "Answer 2 English": resdict["AnswersEng"][1],
            "Answer 3 English": resdict["AnswersEng"][2],
            "Answer 4 English": resdict["AnswersEng"][3],
            "Answer 5 English": resdict["AnswersEng"][4],
            "Answer 1 Sinhala": ConvertToLegacy(resdict["AnswersSin"][0]),
            "Answer 2 Sinhala": ConvertToLegacy(resdict["AnswersSin"][1]),
            "Answer 3 Sinhala": ConvertToLegacy(resdict["AnswersSin"][2]),
            "Answer 4 Sinhala": ConvertToLegacy(resdict["AnswersSin"][3]),
            "Answer 5 Sinhala": ConvertToLegacy(resdict["AnswersSin"][4]),

            "Explanation 1 English": resdict["ExplEng"][0],
            "Explanation 2 English": resdict["ExplEng"][1],
            "Explanation 3 English": resdict["ExplEng"][2],
            "Explanation 4 English": resdict["ExplEng"][3],
            "Explanation 5 English": resdict["ExplEng"][4],
            "Explanation 1 Sinhala": ConvertToLegacy(resdict["ExplSin"][0]),
            "Explanation 2 Sinhala": ConvertToLegacy(resdict["ExplSin"][1]),
            "Explanation 3 Sinhala": ConvertToLegacy(resdict["ExplSin"][2]),
            "Explanation 4 Sinhala": ConvertToLegacy(resdict["ExplSin"][3]),
            "Explanation 5 Sinhala": ConvertToLegacy(resdict["ExplSin"][4]),

            "QNum": str(resdict["AnsNo"]),
        }

    if qtype == "statement":
        return {
            "Question English": resdict["QEng"],
            "Question Sinhala": ConvertToLegacy(resdict["QSin"]),

            "StateAEng": resdict["StatementsEng"][0],
            "StateBEng": resdict["StatementsEng"][1],
            "StateCEng": resdict["StatementsEng"][2],
            "StateASin": ConvertToLegacy(resdict["StatementsSin"][0]),
            "StateBSin": ConvertToLegacy(resdict["StatementsSin"][1]),
            "StateCSin": ConvertToLegacy(resdict["StatementsSin"][2]),

            "Answer 1 English": resdict["AnswersEng"][0],
            "Answer 2 English": resdict["AnswersEng"][1],
            "Answer 3 English": resdict["AnswersEng"][2],
            "Answer 4 English": resdict["AnswersEng"][3],
            "Answer 5 English": resdict["AnswersEng"][4],
            "Answer 1 Sinhala": ConvertToLegacy(resdict["AnswersSin"][0]),
            "Answer 2 Sinhala": ConvertToLegacy(resdict["AnswersSin"][1]),
            "Answer 3 Sinhala": ConvertToLegacy(resdict["AnswersSin"][2]),
            "Answer 4 Sinhala": ConvertToLegacy(resdict["AnswersSin"][3]),
            "Answer 5 Sinhala": ConvertToLegacy(resdict["AnswersSin"][4]),

            "ExplAEng": resdict["ExplEng"][0],
            "ExplBEng": resdict["ExplEng"][1],
            "ExplCEng": resdict["ExplEng"][2],
            "ExplASin": ConvertToLegacy(resdict["ExplSin"][0]),
            "ExplBSin": ConvertToLegacy(resdict["ExplSin"][1]),
            "ExplCSin": ConvertToLegacy(resdict["ExplSin"][2]),

            "QNum": str(resdict["AnsNo"]),
        }

    if qtype == "code":
        return {
            "Question English": resdict["QEng"],
            "Question Sinhala": ConvertToLegacy(resdict["QSin"]),

            "codelines": resdict["Code"],

            "Answer 1 English": resdict["AnswersEng"][0],
            "Answer 2 English": resdict["AnswersEng"][1],
            "Answer 3 English": resdict["AnswersEng"][2],
            "Answer 4 English": resdict["AnswersEng"][3],
            "Answer 5 English": resdict["AnswersEng"][4],
            "Answer 1 Sinhala": ConvertToLegacy(resdict["AnswersSin"][0]),
            "Answer 2 Sinhala": ConvertToLegacy(resdict["AnswersSin"][1]),
            "Answer 3 Sinhala": ConvertToLegacy(resdict["AnswersSin"][2]),
            "Answer 4 Sinhala": ConvertToLegacy(resdict["AnswersSin"][3]),
            "Answer 5 Sinhala": ConvertToLegacy(resdict["AnswersSin"][4]),

            "ExplanationEnglish": resdict["ExplEng"],
            "ExplanationSinhala": ConvertToLegacy(resdict["ExplSin"]),

            "QNum": str(resdict["AnsNo"]),
        }

    raise ValueError(f"Unknown question type: {qtype}")


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.route("/api/generate", methods=["POST"])
def generate():
    payload = request.get_json(silent=True) or {}

    topic = (payload.get("topic") or "").strip()
    model = (payload.get("model") or "").strip()
    qtype = (payload.get("qtype") or "normal").strip().lower()

    if not topic:
        return jsonify({"error": "Please enter a question topic."}), 400
    if not model:
        return jsonify({"error": "Please enter a Gemini model name."}), 400
    if qtype not in VALID_TYPES:
        return jsonify({"error": f"Unknown question type '{qtype}'."}), 400

    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return jsonify({
            "error": "The server is missing a GEMINI_API_KEY environment variable. "
                     "Add it in your Vercel project settings and redeploy."
        }), 500

    try:
        prompt = PROMPT_PATH.read_text(encoding="utf-8")
    except OSError:
        return jsonify({"error": "The prompt file is missing on the server."}), 500

    question_type_and_topic = json.dumps({"topic": topic, "type": qtype}, ensure_ascii=False)

    try:
        client = genai.Client(api_key=api_key)
        response_text = client.models.generate_content(
            model=model,
            contents=prompt + question_type_and_topic,
        ).text
    except Exception as exc:  # noqa: BLE001 - surface the provider's error to the user
        return jsonify({"error": f"The Gemini API request failed: {exc}"}), 502

    try:
        resdict = extract_json(response_text or "")
    except (ValueError, json.JSONDecodeError):
        return jsonify({
            "error": "The model didn't return valid JSON, so no document could be built. "
                     "Try again, or try a different model."
        }), 502

    result_qtype = str(resdict.get("QType", qtype)).strip().lower()
    if result_qtype not in VALID_TYPES:
        return jsonify({"error": f"The model returned an unknown QType: {result_qtype!r}."}), 502

    try:
        replacements = build_replacements(resdict, result_qtype)
        doc = Document(TEMPLATE_FILES[result_qtype])
        find_and_replace(doc, replacements)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
    except (KeyError, IndexError) as exc:
        return jsonify({
            "error": f"The model's response was missing an expected field ({exc}). Try again."
        }), 502
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": f"Failed to build the document: {exc}"}), 500

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_topic = re.sub(r"[^A-Za-z0-9]+", "_", topic).strip("_")[:40] or "MCQ"
    filename = f"{safe_topic}_{result_qtype}_{timestamp}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/generate", methods=["GET"])
def generate_info():
    # Friendly response if someone opens the endpoint directly in a browser.
    return jsonify({"error": "Use POST with a JSON body to generate a question."}), 405
